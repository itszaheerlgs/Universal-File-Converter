"""
Universal File Converter
Supports: PDF↔Word, Images, Excel↔CSV, Text↔PDF, Markdown↔HTML, and more
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import threading
import os
import sys

# ── colour palette ──────────────────────────────────────────────────────────
BG        = "#1e1e2e"
PANEL     = "#2a2a3e"
ACCENT    = "#7c6af7"
ACCENT2   = "#5e9cf7"
SUCCESS   = "#4ade80"
WARNING   = "#f59e0b"
DANGER    = "#f87171"
TEXT      = "#e2e8f0"
SUBTEXT   = "#94a3b8"
CARD      = "#313149"
HOVER     = "#3a3a55"
WHITE     = "#ffffff"

FONT_H    = ("Segoe UI", 14, "bold")
FONT_M    = ("Segoe UI", 11)
FONT_S    = ("Segoe UI", 10)
FONT_XS   = ("Segoe UI", 9)

# ── conversion catalogue ────────────────────────────────────────────────────
CONVERTERS = [
    # (display name, category, from_ext list, to_ext list, func_name)
    ("PDF → Word (.docx)",       "📄 Documents",  [".pdf"],  [".docx"], "pdf_to_word"),
    ("Word → PDF",               "📄 Documents",  [".docx",".doc"], [".pdf"], "word_to_pdf"),
    ("PDF → Text (.txt)",        "📄 Documents",  [".pdf"],  [".txt"],  "pdf_to_text"),
    ("Text → PDF",               "📄 Documents",  [".txt"],  [".pdf"],  "text_to_pdf"),
    ("Markdown → HTML",          "📄 Documents",  [".md",".markdown"], [".html"], "md_to_html"),
    ("HTML → Text",              "📄 Documents",  [".html",".htm"], [".txt"], "html_to_text"),
    ("Excel → CSV",              "📊 Spreadsheets", [".xlsx",".xls"], [".csv"], "excel_to_csv"),
    ("CSV → Excel",              "📊 Spreadsheets", [".csv"], [".xlsx"], "csv_to_excel"),
    ("Image → PDF",              "🖼 Images",     [".jpg",".jpeg",".png",".bmp",".tiff",".webp"], [".pdf"], "image_to_pdf"),
    ("PDF → Images (PNG)",       "🖼 Images",     [".pdf"],  [".png"],  "pdf_to_images"),
    ("Image Format Convert",     "🖼 Images",     [".jpg",".jpeg",".png",".bmp",".tiff",".webp",".gif"], [".jpg",".png",".bmp",".tiff",".webp",".gif"], "convert_image"),
    ("Merge PDFs",               "🔧 PDF Tools",  [".pdf"],  [".pdf"],  "merge_pdfs"),
    ("Split PDF (per page)",     "🔧 PDF Tools",  [".pdf"],  [".pdf"],  "split_pdf"),
    ("Rotate PDF Pages",         "🔧 PDF Tools",  [".pdf"],  [".pdf"],  "rotate_pdf"),
    ("Encrypt PDF",              "🔧 PDF Tools",  [".pdf"],  [".pdf"],  "encrypt_pdf"),
    ("Word → Text",              "📄 Documents",  [".docx",".doc"], [".txt"], "word_to_text"),
    ("JSON → CSV",               "📊 Spreadsheets", [".json"], [".csv"], "json_to_csv"),
    ("CSV → JSON",               "📊 Spreadsheets", [".csv"], [".json"], "csv_to_json"),
]

# ── individual conversion functions ─────────────────────────────────────────

def pdf_to_word(src, dst, **kw):
    from pdf2docx import Converter
    cv = Converter(src)
    cv.convert(dst, start=0, end=None)
    cv.close()

def word_to_pdf(src, dst, **kw):
    # Uses LibreOffice if available, else python-docx + reportlab fallback
    import subprocess, shutil
    if shutil.which("soffice"):
        out_dir = os.path.dirname(dst) or "."
        subprocess.run(["soffice","--headless","--convert-to","pdf",
                        "--outdir", out_dir, src], check=True)
        # soffice names output as <basename>.pdf
        base = os.path.splitext(os.path.basename(src))[0] + ".pdf"
        generated = os.path.join(out_dir, base)
        if generated != dst:
            os.rename(generated, dst)
    else:
        # Fallback: extract text → pdf via reportlab
        from docx import Document as DocxDoc
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        doc = DocxDoc(src)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        _text_to_pdf_content(full_text, dst)

def _text_to_pdf_content(text, dst):
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    styles = getSampleStyleSheet()
    pdf = SimpleDocTemplate(dst, pagesize=letter)
    story = []
    for line in text.split("\n"):
        story.append(Paragraph(line.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;") or " ", styles["Normal"]))
        story.append(Spacer(1, 4))
    pdf.build(story)

def pdf_to_text(src, dst, **kw):
    import pdfplumber
    with pdfplumber.open(src) as pdf:
        text = "\n\n".join(page.extract_text() or "" for page in pdf.pages)
    with open(dst, "w", encoding="utf-8") as f:
        f.write(text)

def text_to_pdf(src, dst, **kw):
    with open(src, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    _text_to_pdf_content(text, dst)

def md_to_html(src, dst, **kw):
    import markdown2
    with open(src, "r", encoding="utf-8") as f:
        md = f.read()
    html = markdown2.markdown(md, extras=["tables","fenced-code-blocks","strike"])
    full = f"""<!DOCTYPE html><html><head><meta charset='utf-8'>
<style>body{{font-family:sans-serif;max-width:860px;margin:auto;padding:2em;line-height:1.6}}
pre{{background:#f4f4f4;padding:1em;border-radius:4px;overflow-x:auto}}
code{{background:#f4f4f4;padding:.2em .4em;border-radius:3px}}
table{{border-collapse:collapse;width:100%}}td,th{{border:1px solid #ccc;padding:.5em}}</style>
</head><body>{html}</body></html>"""
    with open(dst, "w", encoding="utf-8") as f:
        f.write(full)

def html_to_text(src, dst, **kw):
    from html.parser import HTMLParser
    class _P(HTMLParser):
        def __init__(self):
            super().__init__(); self.parts=[]
        def handle_data(self, d): self.parts.append(d)
    with open(src,"r",encoding="utf-8",errors="replace") as f:
        raw=f.read()
    p=_P(); p.feed(raw)
    with open(dst,"w",encoding="utf-8") as f:
        f.write("".join(p.parts))

def excel_to_csv(src, dst, **kw):
    import openpyxl, csv
    wb = openpyxl.load_workbook(src, data_only=True)
    ws = wb.active
    with open(dst,"w",newline="",encoding="utf-8") as f:
        w = csv.writer(f)
        for row in ws.iter_rows(values_only=True):
            w.writerow([c if c is not None else "" for c in row])

def csv_to_excel(src, dst, **kw):
    import csv, openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    with open(src,"r",encoding="utf-8",errors="replace") as f:
        for row in csv.reader(f):
            ws.append(row)
    wb.save(dst)

def image_to_pdf(src, dst, **kw):
    from PIL import Image
    img = Image.open(src).convert("RGB")
    img.save(dst, "PDF", resolution=150)

def pdf_to_images(src, dst_pattern, **kw):
    """dst_pattern = /path/to/output_dir  (we create page_001.png etc)"""
    from PIL import Image
    try:
        import pypdfium2 as pdfium
        doc = pdfium.PdfDocument(src)
        base = os.path.splitext(os.path.basename(src))[0]
        out_dir = dst_pattern
        os.makedirs(out_dir, exist_ok=True)
        for i, page in enumerate(doc):
            bm = page.render(scale=2)
            img = bm.to_pil()
            img.save(os.path.join(out_dir, f"{base}_page_{i+1:03d}.png"))
        return out_dir
    except ImportError:
        # fallback: pdf2image / poppler
        from pdf2image import convert_from_path
        base = os.path.splitext(os.path.basename(src))[0]
        out_dir = dst_pattern
        os.makedirs(out_dir, exist_ok=True)
        imgs = convert_from_path(src, dpi=150)
        for i, img in enumerate(imgs):
            img.save(os.path.join(out_dir, f"{base}_page_{i+1:03d}.png"))
        return out_dir

def convert_image(src, dst, **kw):
    from PIL import Image
    img = Image.open(src)
    ext = os.path.splitext(dst)[1].lower()
    if ext in (".jpg", ".jpeg"):
        img = img.convert("RGB")
    img.save(dst)

def merge_pdfs(sources, dst, **kw):
    from pypdf import PdfWriter
    writer = PdfWriter()
    for s in sources:
        from pypdf import PdfReader
        reader = PdfReader(s)
        for page in reader.pages:
            writer.add_page(page)
    with open(dst,"wb") as f:
        writer.write(f)

def split_pdf(src, dst_dir, **kw):
    from pypdf import PdfReader, PdfWriter
    reader = PdfReader(src)
    base = os.path.splitext(os.path.basename(src))[0]
    os.makedirs(dst_dir, exist_ok=True)
    for i, page in enumerate(reader.pages):
        w = PdfWriter()
        w.add_page(page)
        with open(os.path.join(dst_dir, f"{base}_page_{i+1:03d}.pdf"),"wb") as f:
            w.write(f)

def rotate_pdf(src, dst, degrees=90, **kw):
    from pypdf import PdfReader, PdfWriter
    reader = PdfReader(src)
    writer = PdfWriter()
    for page in reader.pages:
        page.rotate(int(degrees))
        writer.add_page(page)
    with open(dst,"wb") as f:
        writer.write(f)

def encrypt_pdf(src, dst, password="password", **kw):
    from pypdf import PdfReader, PdfWriter
    reader = PdfReader(src)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt(password)
    with open(dst,"wb") as f:
        writer.write(f)

def word_to_text(src, dst, **kw):
    from docx import Document as DocxDoc
    doc = DocxDoc(src)
    with open(dst,"w",encoding="utf-8") as f:
        for p in doc.paragraphs:
            f.write(p.text + "\n")

def json_to_csv(src, dst, **kw):
    import json, csv
    with open(src,"r",encoding="utf-8") as f:
        data = json.load(f)
    if isinstance(data, dict):
        data = [data]
    if not data:
        open(dst,"w").close(); return
    keys = list(data[0].keys())
    with open(dst,"w",newline="",encoding="utf-8") as f:
        w = csv.DictWriter(f, fieldnames=keys)
        w.writeheader()
        w.writerows(data)

def csv_to_json(src, dst, **kw):
    import csv, json
    rows = []
    with open(src,"r",encoding="utf-8",errors="replace") as f:
        for row in csv.DictReader(f):
            rows.append(row)
    with open(dst,"w",encoding="utf-8") as f:
        json.dump(rows, f, indent=2, ensure_ascii=False)

# ── main app ────────────────────────────────────────────────────────────────

class ConverterApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Universal File Converter - I.T Ngani | Dether/Zaheer S. Lagos")
        self.geometry("1000x700")
        self.minsize(800, 560)
        self.configure(bg=BG)
        self.resizable(True, True)

        # state
        self.selected_converter = tk.StringVar()
        self.input_files  = []          # list of paths
        self.output_path  = tk.StringVar()
        self.status_msg   = tk.StringVar(value="Select a converter to begin.")
        self.progress_var = tk.DoubleVar()
        # extra options
        self.extra_degrees   = tk.StringVar(value="90")
        self.extra_password  = tk.StringVar(value="password")
        self.extra_img_fmt   = tk.StringVar(value=".png")
        self.extra_out_fmt   = tk.StringVar(value=".jpg")

        self._build_ui()
        self._apply_styles()

    # ── UI construction ──────────────────────────────────────────────────────

    def _build_ui(self):
        # ── left sidebar ──
        sidebar = tk.Frame(self, bg=PANEL, width=260)
        sidebar.pack(side="left", fill="y")
        sidebar.pack_propagate(False)

        tk.Label(sidebar, text="🔄 Converter", font=("Segoe UI",15,"bold"),
                 bg=PANEL, fg=WHITE).pack(pady=(18,4), padx=18, anchor="w")
        tk.Label(sidebar, text="Pick a conversion type",
                 font=FONT_XS, bg=PANEL, fg=SUBTEXT).pack(padx=18, anchor="w")

        # search
        search_frame = tk.Frame(sidebar, bg=PANEL)
        search_frame.pack(fill="x", padx=12, pady=(12,4))
        self.search_var = tk.StringVar()
        self.search_var.trace("w", lambda *_: self._filter_list())
        tk.Entry(search_frame, textvariable=self.search_var,
                 font=FONT_S, bg=CARD, fg=TEXT, insertbackground=TEXT,
                 relief="flat", bd=4).pack(fill="x")

        # converter list
        list_frame = tk.Frame(sidebar, bg=PANEL)
        list_frame.pack(fill="both", expand=True, padx=6, pady=4)

        self.listbox = tk.Listbox(list_frame, font=FONT_S, bg=PANEL,
                                  fg=TEXT, selectbackground=ACCENT,
                                  selectforeground=WHITE, activestyle="none",
                                  relief="flat", bd=0, highlightthickness=0,
                                  cursor="hand2")
        sb = tk.Scrollbar(list_frame, orient="vertical",
                          command=self.listbox.yview)
        self.listbox.configure(yscrollcommand=sb.set)
        self.listbox.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")
        self.listbox.bind("<<ListboxSelect>>", self._on_select)

        self._all_entries = []
        self._populate_list()

        # ── right main area ──
        main = tk.Frame(self, bg=BG)
        main.pack(side="right", fill="both", expand=True)

        # header
        hdr = tk.Frame(main, bg=BG)
        hdr.pack(fill="x", padx=24, pady=(20,0))
        self.title_lbl = tk.Label(hdr, text="Choose a Converter →",
                                  font=("Segoe UI",17,"bold"), bg=BG, fg=WHITE)
        self.title_lbl.pack(side="left")

        # card
        card = tk.Frame(main, bg=CARD, bd=0)
        card.pack(fill="both", expand=True, padx=24, pady=14)

        # ── input row ──
        r_in = tk.Frame(card, bg=CARD)
        r_in.pack(fill="x", padx=18, pady=(18,6))
        tk.Label(r_in, text="Input File(s)", font=FONT_M, bg=CARD, fg=TEXT,
                 width=12, anchor="w").pack(side="left")
        self.in_entry = tk.Entry(r_in, font=FONT_S, bg=PANEL, fg=TEXT,
                                 insertbackground=TEXT, relief="flat", bd=4)
        self.in_entry.pack(side="left", fill="x", expand=True, padx=(6,6))
        tk.Button(r_in, text="Browse", font=FONT_S, bg=ACCENT, fg=WHITE,
                  relief="flat", bd=0, padx=12, pady=4, cursor="hand2",
                  command=self._browse_input).pack(side="left")

        # ── output row ──
        r_out = tk.Frame(card, bg=CARD)
        r_out.pack(fill="x", padx=18, pady=6)
        tk.Label(r_out, text="Output", font=FONT_M, bg=CARD, fg=TEXT,
                 width=12, anchor="w").pack(side="left")
        self.out_entry = tk.Entry(r_out, textvariable=self.output_path,
                                  font=FONT_S, bg=PANEL, fg=TEXT,
                                  insertbackground=TEXT, relief="flat", bd=4)
        self.out_entry.pack(side="left", fill="x", expand=True, padx=(6,6))
        tk.Button(r_out, text="Browse", font=FONT_S, bg=ACCENT2, fg=WHITE,
                  relief="flat", bd=0, padx=12, pady=4, cursor="hand2",
                  command=self._browse_output).pack(side="left")

        # ── options frame (dynamic) ──
        self.opts_frame = tk.Frame(card, bg=CARD)
        self.opts_frame.pack(fill="x", padx=18, pady=6)

        # ── convert button ──
        btn_row = tk.Frame(card, bg=CARD)
        btn_row.pack(pady=10)
        self.conv_btn = tk.Button(btn_row, text="⚡  Convert Now",
                                  font=("Segoe UI",12,"bold"),
                                  bg=ACCENT, fg=WHITE, relief="flat", bd=0,
                                  padx=28, pady=10, cursor="hand2",
                                  command=self._run_conversion)
        self.conv_btn.pack()

        # ── progress ──
        prog_frame = tk.Frame(card, bg=CARD)
        prog_frame.pack(fill="x", padx=18, pady=(4,0))
        self.pbar = ttk.Progressbar(prog_frame, variable=self.progress_var,
                                    maximum=100, mode="determinate",
                                    style="Accent.Horizontal.TProgressbar")
        self.pbar.pack(fill="x")

        # ── log ──
        log_frame = tk.Frame(card, bg=CARD)
        log_frame.pack(fill="both", expand=True, padx=18, pady=(10,18))
        tk.Label(log_frame, text="Log", font=FONT_XS, bg=CARD,
                 fg=SUBTEXT).pack(anchor="w")
        self.log = tk.Text(log_frame, font=("Consolas",9), bg=PANEL, fg=TEXT,
                           insertbackground=TEXT, relief="flat", bd=0,
                           state="disabled", wrap="word", height=8)
        log_sb = tk.Scrollbar(log_frame, orient="vertical",
                              command=self.log.yview)
        self.log.configure(yscrollcommand=log_sb.set)
        self.log.pack(side="left", fill="both", expand=True)
        log_sb.pack(side="right", fill="y")

        # ── status bar ──
        bar = tk.Frame(self, bg=PANEL, height=30)
        bar.pack(side="bottom", fill="x")
        tk.Label(bar, textvariable=self.status_msg,
                 font=FONT_XS, bg=PANEL, fg=SUBTEXT).pack(side="left", padx=12)
        tk.Label(bar, text="✦  Programmed by Dether / Zaheer S. Lagos  ✦",
                 font=("Segoe UI", 9, "italic"), bg=PANEL, fg=ACCENT).pack(side="right", padx=14)

    def _apply_styles(self):
        s = ttk.Style(self)
        s.theme_use("clam")
        s.configure("Accent.Horizontal.TProgressbar",
                    troughcolor=PANEL, background=ACCENT,
                    borderwidth=0, thickness=8)

    def _populate_list(self, query=""):
        self.listbox.delete(0, "end")
        self._all_entries.clear()
        current_cat = None
        for (name, cat, fin, fout, func) in CONVERTERS:
            if query and query.lower() not in name.lower():
                continue
            if cat != current_cat:
                if not query:
                    self.listbox.insert("end", f"  {cat}")
                    self.listbox.itemconfig("end", fg=ACCENT, selectbackground=PANEL,
                                           selectforeground=ACCENT)
                    self._all_entries.append(None)   # separator
                current_cat = cat
            self.listbox.insert("end", f"     {name}")
            self._all_entries.append((name, cat, fin, fout, func))

    def _filter_list(self):
        self._populate_list(self.search_var.get())

    def _on_select(self, _=None):
        sel = self.listbox.curselection()
        if not sel:
            return
        idx = sel[0]
        entry = self._all_entries[idx]
        if entry is None:   # category header
            return
        name, cat, fin, fout, func = entry
        self.selected_converter.set(func)
        self.title_lbl.config(text=name)
        self.status_msg.set(f"Selected: {name}")
        self.input_files.clear()
        self.in_entry.delete(0, "end")
        self.output_path.set("")
        self._refresh_options(func)

    def _refresh_options(self, func):
        for w in self.opts_frame.winfo_children():
            w.destroy()

        def lbl(text):
            tk.Label(self.opts_frame, text=text, font=FONT_S,
                     bg=CARD, fg=SUBTEXT, width=18, anchor="w").pack(side="left")

        if func == "rotate_pdf":
            lbl("Rotation degrees:")
            tk.Entry(self.opts_frame, textvariable=self.extra_degrees,
                     font=FONT_S, bg=PANEL, fg=TEXT, insertbackground=TEXT,
                     relief="flat", bd=4, width=8).pack(side="left")

        elif func == "encrypt_pdf":
            lbl("Password:")
            tk.Entry(self.opts_frame, textvariable=self.extra_password,
                     font=FONT_S, bg=PANEL, fg=TEXT, insertbackground=TEXT,
                     relief="flat", bd=4, width=20, show="*").pack(side="left")

        elif func == "convert_image":
            lbl("Output format:")
            fmt_cb = ttk.Combobox(self.opts_frame, textvariable=self.extra_out_fmt,
                                  values=[".jpg",".png",".bmp",".tiff",".webp",".gif"],
                                  width=8, state="readonly", font=FONT_S)
            fmt_cb.pack(side="left")

    # ── file browsing ────────────────────────────────────────────────────────

    def _browse_input(self):
        func = self.selected_converter.get()
        if not func:
            messagebox.showinfo("Select Converter", "Please choose a converter first.")
            return

        # special multi-file for merge
        if func == "merge_pdfs":
            files = filedialog.askopenfilenames(
                title="Select PDF files to merge",
                filetypes=[("PDF files","*.pdf")])
            if files:
                self.input_files = list(files)
                self.in_entry.delete(0,"end")
                self.in_entry.insert(0, "; ".join(os.path.basename(f) for f in files))
                self._auto_output()
            return

        # find accepted extensions
        entry = next((e for e in self._all_entries if e and e[4]==func), None)
        if not entry:
            return
        _, _, fin, _, _ = entry
        ftypes = [(f"Files ({', '.join(fin)})", " ".join(f"*{e}" for e in fin)),
                  ("All files","*.*")]
        path = filedialog.askopenfilename(title="Select input file", filetypes=ftypes)
        if path:
            self.input_files = [path]
            self.in_entry.delete(0,"end")
            self.in_entry.insert(0, path)
            self._auto_output()

    def _auto_output(self):
        func = self.selected_converter.get()
        if not self.input_files:
            return
        src = self.input_files[0]
        base = os.path.splitext(src)[0]
        entry = next((e for e in self._all_entries if e and e[4]==func), None)
        if not entry:
            return
        _, _, _, fout, _ = entry

        # special cases with directory output
        if func in ("pdf_to_images","split_pdf"):
            out = base + "_output"
            self.output_path.set(out)
            return
        if func == "convert_image":
            out = base + self.extra_out_fmt.get()
        elif func == "merge_pdfs":
            out = os.path.join(os.path.dirname(src), "merged.pdf")
        else:
            out = base + fout[0]
        self.output_path.set(out)

    def _browse_output(self):
        func = self.selected_converter.get()
        entry = next((e for e in self._all_entries if e and e[4]==func), None)
        if not entry:
            return
        _, _, _, fout, _ = entry

        if func in ("pdf_to_images","split_pdf"):
            path = filedialog.askdirectory(title="Select output directory")
        else:
            ext = self.extra_out_fmt.get() if func=="convert_image" else fout[0]
            ftypes = [(f"Files (*{ext})", f"*{ext}"), ("All files","*.*")]
            path = filedialog.asksaveasfilename(
                defaultextension=ext, filetypes=ftypes,
                title="Save output as")
        if path:
            self.output_path.set(path)

    # ── conversion runner ────────────────────────────────────────────────────

    def _log(self, msg, color=None):
        self.log.config(state="normal")
        tag = None
        if color:
            tag = f"col_{color}"
            self.log.tag_configure(tag, foreground=color)
        self.log.insert("end", msg+"\n", tag or "")
        self.log.see("end")
        self.log.config(state="disabled")

    def _run_conversion(self):
        func = self.selected_converter.get()
        if not func:
            messagebox.showwarning("No Converter","Please choose a converter first.")
            return
        if not self.input_files:
            messagebox.showwarning("No Input","Please select an input file.")
            return
        if not self.output_path.get():
            messagebox.showwarning("No Output","Please set an output path.")
            return

        self.conv_btn.config(state="disabled", text="Converting…")
        self.progress_var.set(0)
        threading.Thread(target=self._do_convert, daemon=True).start()

    def _do_convert(self):
        func  = self.selected_converter.get()
        src   = self.input_files[0]
        dst   = self.output_path.get()
        srcs  = self.input_files

        try:
            self._log(f"▶ Starting: {func}", ACCENT)
            self.progress_var.set(10)

            fn_map = {
                "pdf_to_word":   lambda: pdf_to_word(src, dst),
                "word_to_pdf":   lambda: word_to_pdf(src, dst),
                "pdf_to_text":   lambda: pdf_to_text(src, dst),
                "text_to_pdf":   lambda: text_to_pdf(src, dst),
                "md_to_html":    lambda: md_to_html(src, dst),
                "html_to_text":  lambda: html_to_text(src, dst),
                "excel_to_csv":  lambda: excel_to_csv(src, dst),
                "csv_to_excel":  lambda: csv_to_excel(src, dst),
                "image_to_pdf":  lambda: image_to_pdf(src, dst),
                "pdf_to_images": lambda: pdf_to_images(src, dst),
                "convert_image": lambda: convert_image(src, dst),
                "merge_pdfs":    lambda: merge_pdfs(srcs, dst),
                "split_pdf":     lambda: split_pdf(src, dst),
                "rotate_pdf":    lambda: rotate_pdf(src, dst, self.extra_degrees.get()),
                "encrypt_pdf":   lambda: encrypt_pdf(src, dst, self.extra_password.get()),
                "word_to_text":  lambda: word_to_text(src, dst),
                "json_to_csv":   lambda: json_to_csv(src, dst),
                "csv_to_json":   lambda: csv_to_json(src, dst),
            }

            self.progress_var.set(30)
            self._log(f"  Input : {src}")
            self._log(f"  Output: {dst}")
            fn_map[func]()
            self.progress_var.set(100)
            self._log(f"✅ Done! Saved to: {dst}", SUCCESS)
            self.status_msg.set("✅ Conversion complete!")
            messagebox.showinfo("Success", f"Conversion complete!\n\n{dst}")

        except Exception as e:
            self.progress_var.set(0)
            self._log(f"❌ Error: {e}", DANGER)
            self.status_msg.set(f"Error: {e}")
            messagebox.showerror("Conversion Failed", str(e))
        finally:
            self.conv_btn.config(state="normal", text="⚡  Convert Now")

# ── entry point ─────────────────────────────────────────────────────────────

if __name__ == "__main__":
    app = ConverterApp()
    app.mainloop()